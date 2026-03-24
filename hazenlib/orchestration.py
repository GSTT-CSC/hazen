"""Orchestration Module for performing multiple tasks."""

# ruff: noqa: T201

from __future__ import annotations

import datetime
import hashlib

# Type Checking
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    # Python imports
    from collections.abc import Sequence

    # Local imports
    from hazenlib.HazenTask import HazenTask

# Python imports
import importlib
import logging
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import TypeVar

# Module imports
import yaml
from docx import Document
from docx.shared import Inches
from packaging import version as packaging_version
from packaging.specifiers import InvalidSpecifier, SpecifierSet
from pydicom import dcmread

# Local imports
import hazenlib.tasks  # noqa: F401
from hazenlib._version import __version__
from hazenlib.ACRObject import ACRObject
from hazenlib.exceptions import (
    UnknownAcquisitionTypeError,
    UnknownTaskNameError,
)
from hazenlib.types import Measurement, PhantomType, Result, TaskMetadata
from hazenlib.utils import get_dicom_files, wait_on_parallel_results

logger = logging.getLogger(__name__)


def init_task(
    selected_task: str,
    files: list[str],
    report: bool = False,
    report_dir: str | None = None,
    **kwargs,
) -> HazenTask:
    """Initialise object of the correct HazenTask class.

    Args:
        selected_task (string): name of task script/module to load
        files (list): list of filepaths to DICOM images
        report (bool): whether to generate report images
        report_dir (string): path to folder to save report images to
        kwargs: any other key word arguments

    Returns:
        an object of the specified HazenTask class

    """
    try:
        meta = TASK_REGISTRY[selected_task]
    except KeyError as err:
        msg = f"Unknown task: {selected_task}"
        logger.exception(
            "%s. Supported tasks are:\n%s",
            msg,
            "\n\t".join(TASK_REGISTRY),
        )
        raise ValueError(msg) from err

    # Import module
    task_module = importlib.import_module(f"hazenlib.tasks.{meta.module_name}")

    # Get explicit class
    try:
        task_class = getattr(task_module, meta.class_name)
    except AttributeError as err:
        msg = f"Module {meta.module_name} has no class '{meta.class_name}'"
        raise ImportError(msg) from err

    return task_class(
        input_data=files,
        report=report,
        report_dir=report_dir,
        **kwargs,
    )


class AcquisitionType(Enum):
    """Supported Acquisition Types."""

    ACR_T1 = "ACR T1"
    ACR_T2 = "ACR T2"
    ACR_SL = "ACR Sagittal Localizer"

    @classmethod
    def from_string(cls, value: str) -> AcquisitionType:
        """Create AcquisitionType from string with fuzzy matching.

        Handles British/American spelling variations (Localiser vs Localizer)
        and case insensitivity.

        Args:
            value: String representation of acquisition type

        Returns:
            AcquisitionType enum value

        Raises:
            UnknownAcquisitionTypeError: If string cannot be matched to
                a known type

        Example:
            >>> AcquisitionType.from_string("t1")
            AcquisitionType.ACR_T1
            >>> AcquisitionType.from_string("sagittal localiser")
            AcquisitionType.ACR_SL

        """
        normalized = value.lower().replace("localiser", "localizer")
        mapping = {
            "t1": cls.ACR_T1,
            "t2": cls.ACR_T2,
            "sagittal localizer": cls.ACR_SL,
            "sagittal": cls.ACR_SL,
            "localizer": cls.ACR_SL,
        }
        if normalized in mapping:
            return mapping[normalized]
        raise UnknownAcquisitionTypeError(value)


@dataclass(frozen=True)
class ProtocolStep:
    """Single protocol step representing the execution of a single task.

    Attributes:
        task_name: Name of the task as registered in TASK_REGISTRY
        acquisition_type: Type of acquisition this step applies to

    """

    task_name: str
    acquisition_type: AcquisitionType

    def __post_init__(self) -> None:
        """Validate that task name exists in the global registry."""
        if self.task_name not in TASK_REGISTRY:
            available = ", ".join(sorted(TASK_REGISTRY.keys()))
            raise UnknownTaskNameError(self.task_name, available)


@dataclass
class Protocol:
    """Orchestrator for collections of tasks.

    Attributes:
        name: Human-readable identifier for this protocol
        steps: Ordered collection of protocol steps to execute

    """

    name: str
    steps: tuple[ProtocolStep, ...] = field(default_factory=tuple)

    @classmethod
    def from_config(cls, config_path: Path) -> Protocol:
        """Load protocol from configuration file."""
        msg = "The class method 'from_config' has not been implemented yet."
        raise NotImplementedError(msg)


class ProtocolResult(Result):
    """Class for the protocol result."""

    def __post_init__(self, _load_metadata: bool) -> None:
        """Initialise the results list."""
        super().__post_init__(_load_metadata)

        # Set initial result to contain
        # protocol information.
        self._results: list[Result] = [
            Result(
                self.task,
                self.desc,
                self.files,
            ),
        ]

    @property
    def measurements(self) -> tuple[Measurement, ...]:
        """Tuple of initial result measurements."""
        return self.results[0].measurements

    def add_measurement(self, measurement: Measurement) -> None:
        """Add a measurement to the initial result."""
        self.results[0].add_measurement(measurement)

    @property
    def report_images(self) -> tuple[str, ...]:
        """Tuple of initial report image locations."""
        return self.results[0].report_images

    def add_report_image(self, image_path: str | Sequence[str]) -> None:
        """Add a report image to the initial result."""
        return self.results[0].add_report_image(image_path)

    @property
    def results(self) -> tuple[Result, ...]:
        """Return an immutable results list."""
        return tuple(self._results)

    def add_result(self, result: Result) -> None:
        """Add a result to the list."""
        self._results.append(result)

    def add_metadata_to_doc(self, doc: Document) -> Document:
        """Add metadata properties to a document."""
        doc.core_properties.author = "Hazen"
        doc.core_properties.created = datetime.datetime.now(tz=datetime.UTC)
        doc.core_properties.comments = (
            "Initial draft of the document automatically generated by"
            f" Hazen (version {self.metadata.version})"
        )
        return doc

    def to_docx(
        self,
        template_path: Path | str | None = None,
        level: str = "all",
    ) -> Document:
        """Generate Word document from aggregated results."""
        if self.results is None:
            msg = "Results cannot be empty"
            logger.error(
                "%s, please make sure the protocols run method"
                " has been called.",
                msg,
            )
            raise ValueError(msg)

        if template_path is not None:
            template_path = Path(template_path)
            if not template_path.exists():
                msg = f"Template file not found: {template_path}"
                logger.error(msg)
                raise FileNotFoundError(msg)
            doc = Document(template_path)
        else:
            doc = Document()

            # Header with protocol info
            doc.add_heading(f"QA Report: {self.task}", 0)

        # Metadata
        doc = self.add_metadata_to_doc(doc)
        # Section per step
        for _result in self.results:
            # Skip the metadata result for the current protocol
            if _result.task == self.task:
                continue
            doc = add_report_table_to_doc(doc, _result, level)

        return doc


def add_report_table_to_doc(
    doc: Document,
    _result: Result,
    level: str,
) -> Document:
    """Add a report table to the word document."""
    doc.add_heading(_result.task, level=1)

    # Filter out specific results ('all', 'final', 'intermediate', etc.)
    result = _result.filtered(level)

    # internal Measurement property -> word heading
    text_mapping = {
        "name": "Name",
        "type": "Type",
        "subtype": "Subtype",
        "description": "Description",
        "value": "Value",
        "unit": "Unit",
    }

    # Results table
    table = doc.add_table(rows=1, cols=len(text_mapping))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for idx, text in enumerate(text_mapping.values()):
        hdr_cells[idx].text = text

    for m in result.measurements:
        row_cells = table.add_row().cells
        for idx, key in enumerate(text_mapping.keys()):
            value = getattr(m, key)
            row_cells[idx].text = str(value) if value is not None else "-"

    # Embed report images if generated
    if level != "final":  # Ignore for final reports.
        for img_path in result.report_images:
            if Path(img_path).exists():
                doc.add_picture(img_path, width=Inches(5.0))
    return doc


T = TypeVar("T")


class ACRLargePhantomProtocol(Protocol):
    """Protocol for ACR Large Phantom."""

    def __init__(self, dirs: list[str | Path], **kwargs: T.kwargs) -> None:
        """Set up the protocol for the ACR Large Phantom."""
        self.name = "ACR Large Phantom"
        self.steps = (
            # Geometric Accuracy.
            ProtocolStep("acr_geometric_accuracy", AcquisitionType.ACR_T1),
            ProtocolStep(
                "acr_sagittal_geometric_accuracy",
                AcquisitionType.ACR_SL,
            ),
            # High Contrast Object Detection.
            ProtocolStep("acr_spatial_resolution", AcquisitionType.ACR_T1),
            ProtocolStep("acr_spatial_resolution", AcquisitionType.ACR_T2),
            # Slice Thickness Accuracy.
            ProtocolStep("acr_slice_thickness", AcquisitionType.ACR_T1),
            ProtocolStep("acr_slice_thickness", AcquisitionType.ACR_T2),
            # Slice Position Accuracy.
            ProtocolStep("acr_slice_position", AcquisitionType.ACR_T1),
            ProtocolStep("acr_slice_position", AcquisitionType.ACR_T2),
            # Image Intensity Uniformity.
            ProtocolStep("acr_uniformity", AcquisitionType.ACR_T1),
            ProtocolStep("acr_uniformity", AcquisitionType.ACR_T2),
            # Percent Signal Ghosting.
            ProtocolStep("acr_ghosting", AcquisitionType.ACR_T1),
            # Low Contrast Object Detectability.
            ProtocolStep(
                "acr_low_contrast_object_detectability",
                AcquisitionType.ACR_T1,
            ),
            ProtocolStep(
                "acr_low_contrast_object_detectability",
                AcquisitionType.ACR_T2,
            ),
            # SNR
            ProtocolStep("acr_snr", AcquisitionType.ACR_T1),
            ProtocolStep("acr_snr", AcquisitionType.ACR_T2),
        )

        kwargs.setdefault("report", False)
        kwargs.setdefault("report_dir", None)
        self.kwargs = kwargs

        required_acquisition_types = {s.acquisition_type for s in self.steps}
        if len(dirs) != (num_aq := len(required_acquisition_types)):
            msg = f"Incorrect number of directories - should be {num_aq}"
            logger.exception("%s but got %i", msg, len(dirs))
            raise ValueError(msg)

        files_list = (get_dicom_files(d) for d in dirs)
        self.file_groups = {}
        for files in files_list:
            acr_obj = ACRObject([dcmread(files[0], stop_before_pixels=False)])
            acquisition_type = AcquisitionType.from_string(
                acr_obj.acquisition_type(strict=True),
            )
            self.file_groups[acquisition_type] = files

        if len(self.file_groups.keys()) != len(required_acquisition_types):
            acquisition_types = set(self.file_groups.keys())
            msg = (
                f"Missing sequences - only {acquisition_types}"
                " were found meaning"
                f" {required_acquisition_types - acquisition_types}"
                " were not present."
            )
            logger.exception(
                "%s The following directories were passed: %s",
                msg,
                dirs,
            )
            raise ValueError(msg)

    def run(self, *, debug: bool = False) -> ProtocolResult:
        """Run the Protocol for each of the steps."""
        results = ProtocolResult(
            self.name,
            ", ".join(
                f"{s.task_name} {s.acquisition_type}" for s in self.steps
            ),
            list(self.file_groups.values()),
        )

        arg_list = [
            (step, self.file_groups, self.kwargs) for step in self.steps
        ]
        parallel_results = wait_on_parallel_results(
            _execute_step,
            arg_list,
            debug=debug,
        )
        for r in parallel_results:
            results.add_result(r)

        return results


def _execute_step(
    step: ProtocolStep,
    file_groups: dict,
    kwargs: T.kwargs,
) -> Result:
    """Encapsulate the work for a single step."""
    task = init_task(
        step.task_name,
        file_groups[step.acquisition_type],
        **kwargs,
    )
    return task.run()


@dataclass(frozen=True)
class JobTaskConfig:
    """Configuration for a specific task within a job.

    Attributes:
        task : Must match the TASK_REGISTRY or the PROTOCOL_REGISTRY key.
        folders : Path to the folders containing the DICOMs.
        overrides : Task specific overrides.

    """

    task: str
    folders: list[Path]
    overrides: dict[str, Any] = field(default_factory=dict)
    _is_protocol: bool = field(init=False, compare=False, repr=False)

    def __post_init__(self) -> None:
        """Initialise the job task configuration."""
        if self.task in TASK_REGISTRY:
            object.__setattr__(self, "_is_protocol", False)
        elif self.task in PROTOCOL_REGISTRY:
            object.__setattr__(self, "_is_protocol", True)
            # Protocol-specific validation
            if self.task == "acr_all" and len(self.folders) != 3:
                msg = (
                    "acr_all expects 3 directories (T1, T2, Sagittal),"
                    f" got {len(self.folders)}"
                )
                raise ValueError(msg)
        else:
            available_tasks = ", ".join(TASK_REGISTRY.keys())
            available_protocols = ",".join(PROTOCOL_REGISTRY.keys())
            raise UnknownTaskNameError(
                self.task,
                f"{available_protocols}, {available_tasks}",
            )

        for folder in self.folders:
            if not Path(folder).exists():
                msg = f"Folder not found: {folder}"
                raise FileNotFoundError(msg)

    @property
    def is_protocol(self) -> bool:
        """Return True if this job references a protocol."""
        return self._is_protocol


@dataclass
class BatchConfig:
    """Configuration for the batch configuration file."""

    version: str
    hazen_version_constraint: str | None
    description: str
    jobs: list[JobTaskConfig]
    output: str
    report_docx: str | None = None
    report_template: str | None = None
    defaults: dict[str, Any] | None = None
    levels: list[str] | tuple[str] = ("final", "all")

    _file: str | Path | None = field(default=None)
    _dry_run: bool = False

    _CURRENT_BATCHCONFIG_VERSION: str = "1.0"

    def to_yaml(self, output: Path | str) -> None:
        """Save the batch configuration object to a yaml file.

        Args:
            output: Path to write the YAML configuration file.

        """

        def resolve_path_as_posix(path: str | Path | None) -> str | None:
            if path is None:
                return None
            return Path(path).absolute().as_posix()

        output = Path(output)
        output.parent.mkdir(parents=True, exist_ok=True)

        # Build the data dictionary in the expected order
        data: dict[str, Any] = {
            "version": self.version,
        }

        data["hazen_version_constraint"] = self.hazen_version_constraint

        data["description"] = self.description

        data["output"] = resolve_path_as_posix(self.output)

        data["levels"] = list(self.levels)

        data["report_docx"] = resolve_path_as_posix(self.report_docx)

        data["report_template"] = resolve_path_as_posix(self.report_template)

        data["defaults"] = {} if self.defaults is None else self.defaults

        data["defaults"]["report"] = self._check_and_set_default_report(
            data["levels"],
            report=data["defaults"].get("report", False),
        )

        # Build jobs list
        jobs_data: list[dict[str, Any]] = []
        for job in self.jobs:
            job_dict: dict[str, Any] = {
                "task": job.task,
                "folders": [resolve_path_as_posix(f) for f in job.folders],
            }
            if job.overrides:
                job_dict["overrides"] = {}
                for k, v in job.overrides.items():
                    if isinstance(v, Path):
                        job_dict["overrides"][k] = resolve_path_as_posix(v)
                    else:
                        job_dict["overrides"][k] = v
            jobs_data.append(job_dict)

        data["jobs"] = jobs_data

        with output.open("w") as f:
            yaml.dump(data, f, default_flow_style=False, sort_keys=False)

    def __post_init__(self) -> None:
        """Log the batch config initial parameters."""
        logger.debug(
            "Performing batch config job from file: %s",
            self._file,
        )

    def run(
        self,
        *,
        dry_run: bool | None = None,
        debug: bool = False,
    ) -> ProtocolResult:
        """Run the batch config tasks."""
        dry_run = self._dry_run if dry_run is None else dry_run

        ############
        # Protocol #
        ############

        protocol_jobs = [j for j in self.jobs if j.is_protocol]
        protocol_arg_list = []
        for job in protocol_jobs:
            kwargs = dict(self.defaults or {})
            kwargs.update(job.overrides)

            folders = job.folders
            protocol_arg_list.append([job.task, folders, kwargs])

        #########
        # Tasks #
        #########

        task_jobs = [j for j in self.jobs if not j.is_protocol]
        task_arg_list = []
        for job in task_jobs:
            kwargs = dict(self.defaults or {})
            kwargs.update(job.overrides)

            files = get_dicom_files(job.folders[0])
            task_arg_list.append([job.task, files, kwargs])

        result_files = [] if self._file is None else [self._file.as_posix()]
        results = ProtocolResult(
            task="Batch Configuration Job",
            desc=self.description,
            files=result_files,
        )

        ###########
        # Dry run #
        ###########

        if dry_run:
            print(f"Configuration valid. Would execute {len(self.jobs)} jobs:")

            # Protocols
            for i, (protocol, folders, kwargs) in enumerate(protocol_arg_list):
                print(
                    f"{i}. Protocol: {protocol}\n"
                    f"\tFolders:     {[f.as_posix() for f in folders]}\n"
                    f"\tParameters:  {kwargs or '(none)'}",
                )

            # Tasks
            for i, (task, files, kwargs) in enumerate(task_arg_list):
                print(
                    f"{i + len(protocol_arg_list)}. Task: {task}\n"
                    f"\tFiles:       {len(files)} DICOM(s)"
                    f" from {Path(files[0]).parent}\n"
                    f"\tParameters:  {kwargs or '(none)'}",
                )
            print("-" * 60 + "\nDry run complete. No Measurements performed.")
            return results

        #######
        # Run #
        #######

        # Protocols
        for job, args in zip(protocol_jobs, protocol_arg_list, strict=True):
            _, dirs, kwargs = args
            protocol = PROTOCOL_REGISTRY[job.task](dirs, **kwargs)
            protocol_results = protocol.run()
            for r in protocol_results.results:
                if r.task == protocol_results.task:
                    continue
                results.add_result(r)

        # Tasks
        parallel_results = wait_on_parallel_results(
            _execute_task,
            task_arg_list,
            debug=debug,
        )

        for r in parallel_results:
            results.add_result(r)

        return results

    @classmethod
    def _migrate_config(
        cls,
        data: dict[str, Any],
        schema_version: str,
    ) -> dict[str, Any]:
        """Migrate the schema version."""
        msg = "Configuration schema migration is not currently supported."
        raise NotImplementedError(msg)

    @classmethod
    def _validate_schema_version(
        cls,
        schema_version: str,
        data: dict[str, Any],
    ) -> dict[str, Any]:
        current_schema = cls._CURRENT_BATCHCONFIG_VERSION
        if packaging_version.parse(schema_version) > packaging_version.parse(
            current_schema,
        ):
            msg = (
                f"Config file schema version {schema_version} is newer than "
                f"supported version {current_schema}"
            )
            logger.error(
                "%s. Please upgrade hazen to use this configuration file.",
                msg,
            )
            raise ValueError(msg)
        if schema_version != current_schema:
            logger.warning(
                "Config file uses schema version %s"
                ", current is %s."
                " Attempting backward-compatible load."
                " Consider updating your config file to the latest schema.",
                schema_version,
                current_schema,
            )
            data = cls._migrate_config(data, schema_version)
        return data

    @staticmethod
    def _validate_hazen_version(constraint_str: str) -> None:
        """Validate that current hazen version satisfies the constraint.

        Raises:
            RuntimeError: If version constraint is not satisfied or invalid.

        """
        try:
            specifier = SpecifierSet(constraint_str)
        except InvalidSpecifier as e:
            msg = (
                f"Invalid hazen_version_constraint '{constraint_str}': {e}. "
                "Use standard semver specifiers like"
                " '>=2.0.0', '~=2.0.0', '==2.0.*' etc."
            )
            raise ValueError(msg) from e

        current_version = packaging_version.parse(__version__)

        if current_version not in specifier:
            msg = (
                "Version constraint mismatch:"
                f" Config requires hazen {constraint_str},"
                f" but you are running {__version__}."
                " Please install a compatible version:\n"
                f"pip install 'hazen{constraint_str}'"
                f" or\nuv tool install 'hazen{constraint_str}'"
            )
            raise RuntimeError(msg)

    @staticmethod
    def _check_and_set_default_report(
        levels: list[str],
        *,
        report: bool,
    ) -> bool:
        if "all" in levels or "intermediate" in levels:
            logger.info(
                "levels = %s requires report=True as a default"
                " so this has been set.",
                levels,
            )
            return True
        return report

    @classmethod
    def from_config(
        cls,
        config_path: str | Path,
        *,
        dry_run: bool = False,
    ) -> BatchConfig:
        """Read a configuration file into a BatchConfig object."""
        config_path = Path(config_path)
        with config_path.open("r") as f:
            data = yaml.safe_load(f)

        if not isinstance(data, dict):
            msg = f"{config_path} is not a properly formatted yaml file"
            raise TypeError(msg)

        constraint_str = data.get("hazen_version_constraint")
        if constraint_str:
            cls._validate_hazen_version(constraint_str)

        schema_version = data.get("version", "1.0")
        data = cls._validate_schema_version(schema_version, data)

        # Resolve paths relative to config file location
        config_dir = config_path.parent

        def resolve_path(path: str | Path) -> Path:
            p = Path(path)
            if p.is_absolute():
                return p
            # Always resolve relative paths with respect to the config file directory
            return (config_dir / p).absolute()

        # Parse jobs
        jobs = []
        for job_data in data.get("jobs", []):
            # Resolve folder paths
            folders = [resolve_path(f) for f in job_data.get("folders", [])]

            # Validate task name against registry
            task_name = job_data["task"]

            jobs.append(
                JobTaskConfig(
                    task=task_name,
                    folders=folders,
                    overrides=job_data.get("overrides", {}),
                ),
            )

        # Handle optional paths
        def resolve_path_as_str(path: Path | None) -> str | None:
            return resolve_path(path).as_posix() if path is not None else None

        report_docx = resolve_path_as_str(data.get("report_docx"))
        report_template = resolve_path_as_str(data.get("report_template"))
        output = resolve_path_as_str(data.get("output"))

        logger.info(
            "Batch config job created from: %s (hash: %s)",
            config_path,
            hashlib.sha256(str(data).encode("utf-8")).hexdigest(),
        )

        levels = data.get("levels", ["final"])
        if isinstance(levels, str):
            levels = [levels]

        defaults = data.get("defaults", {})
        defaults["report"] = cls._check_and_set_default_report(
            levels,
            report=defaults.get("report", False),
        )
        return cls(
            version=data.get("version", "1.0"),
            hazen_version_constraint=data.get("hazen_version_constraint"),
            description=data.get("description", ""),
            output=output,
            jobs=jobs,
            levels=levels,
            report_docx=report_docx,
            report_template=report_template,
            defaults=defaults,
            _file=config_path,
            _dry_run=dry_run,
        )


def _execute_task(
    task: str,
    files: list[str],
    kwargs: dict[str, Any],
) -> Result:
    """Encapsulate the work for a single task."""
    report = kwargs.pop("report", False)
    report_dir = kwargs.pop("report_dir", None)
    task = init_task(
        task,
        files,
        report=report,
        report_dir=report_dir,
        **kwargs,
    )
    return task.run()


##############
# Registries #
##############

PROTOCOL_REGISTRY: dict[str, type[Protocol]] = {
    "acr_all": ACRLargePhantomProtocol,
}

# Note that if changing the TASK_REGISTRY keys you should
# update Protocol to match up with the task registry.
TASK_REGISTRY = {
    # MagNET #
    "snr": TaskMetadata(
        module_name="snr",
        class_name="SNR",
        single_image=True,
        phantom=PhantomType.MAGNET,
    ),
    "ghosting": TaskMetadata(
        module_name="ghosting",
        class_name="Ghosting",
        single_image=True,
        phantom=PhantomType.MAGNET,
    ),
    "uniformity": TaskMetadata(
        module_name="uniformity",
        class_name="Uniformity",
        single_image=True,
        phantom=PhantomType.MAGNET,
    ),
    "spatial_resolution": TaskMetadata(
        module_name="spatial_resolution",
        class_name="SpatialResolution",
        single_image=True,
        phantom=PhantomType.MAGNET,
    ),
    "slice_width": TaskMetadata(
        module_name="slice_width",
        class_name="SliceWidth",
        single_image=True,
        phantom=PhantomType.MAGNET,
    ),
    "slice_position": TaskMetadata(
        module_name="slice_position",
        class_name="SlicePosition",
        single_image=True,
        phantom=PhantomType.MAGNET,
    ),
    "snr_map": TaskMetadata(
        module_name="snr_map",
        class_name="SNRMap",
        single_image=True,
        phantom=PhantomType.MAGNET,
    ),
    # ACR #
    "acr_geometric_accuracy": TaskMetadata(
        module_name="acr_geometric_accuracy",
        class_name="ACRGeometricAccuracy",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_ghosting": TaskMetadata(
        module_name="acr_ghosting",
        class_name="ACRGhosting",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_low_contrast_object_detectability": TaskMetadata(
        module_name="acr_low_contrast_object_detectability",
        class_name="ACRLowContrastObjectDetectability",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_object_detectability": TaskMetadata(
        module_name="acr_object_detectability",
        class_name="ACRObjectDetectability",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_slice_position": TaskMetadata(
        module_name="acr_slice_position",
        class_name="ACRSlicePosition",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_slice_thickness": TaskMetadata(
        module_name="acr_slice_thickness",
        class_name="ACRSliceThickness",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_snr": TaskMetadata(
        module_name="acr_snr",
        class_name="ACRSNR",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_spatial_resolution": TaskMetadata(
        module_name="acr_spatial_resolution",
        class_name="ACRSpatialResolution",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_sagittal_geometric_accuracy": TaskMetadata(
        module_name="acr_sagittal_geometric_accuracy",
        class_name="ACRSagittalGeometricAccuracy",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    "acr_uniformity": TaskMetadata(
        module_name="acr_uniformity",
        class_name="ACRUniformity",
        single_image=False,
        phantom=PhantomType.ACR,
    ),
    # Caliber
    "relaxometry": TaskMetadata(
        module_name="relaxometry",
        class_name="Relaxometry",
        single_image=False,
        phantom=PhantomType.CALIBER,
    ),
}
